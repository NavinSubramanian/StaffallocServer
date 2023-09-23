import random
import os
import pandas as pd
import copy
from docx import Document
from docx.shared import Cm,Inches,Twips
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
import docx2pdf
from PyPDF2 import PdfMerger


# Left shift function to cause no problem
def superlogic(date,exam,rooms,tot,single,girls,mselected_faculty,fselected_faculty):
    day_list = []
    print(girls)

    def leftshift(invigilators):
        shift=3
        for a in range(shift):
            temp=invigilators[0]
            for b in range(len(invigilators)-1):
                invigilators[b]=invigilators[b+1]
            invigilators[b]=temp
        return invigilators

    # Required exam dates here

    dates=date

    # Start of Main function

    def generateResult(roomsPerDay,girls):
        staffIndex = 0
        dates_number = 0
        girl = 0

        #Initializing the count for each staff
        '''
        designation_count = {i:0 for i in invigilators}
        '''

        for j in range(len(roomsPerDay)):

            temp_alls = []
            
            start_ind = 0
            
            session = 'FN' #if j == 0 else 'AN' # set session to FN for first day (internal assessment), AN for other days (model exams)
            print('after2')
            global girlss
            print(date,rooms,single,girls,mselected_faculty,fselected_faculty)
            print('af')
            singles = single 
            girls_rooms={}
            '''Boys rooms dictionary'''
            for i in girls:
                girls_rooms[i]=[]
            girls=set(girls)
            allr=set(rooms)
            byr=allr-girls
            boys_rooms={}
            for i in byr:
                boys_rooms[i]=[]
            print(all_rooms,boys_rooms,girls_rooms,invigilators,b_invi,g_invi)
            all_roomsAlloc = copy.deepcopy(rooms)
            boys_roomAllocation = copy.deepcopy(boys_rooms)
            girls_roomAllocation = copy.deepcopy(girls_rooms)


            #to check the count of staff 'x' designation

            '''            
            while(designation_count[invigilators[staffIndex]] == designation[invigilators[staffIndex]] and staffIndex != totalTeachers):
                staffIndex++
            '''

            if staffIndex == len(b_invi):
                leftshift(invigilators)
                staffIndex = 0

            for room in all_roomsAlloc:
                
                if room in girls_roomAllocation and room not in singles and g_invi[girl] not in temp_alls:
                    girls_roomAllocation[room].append(g_invi[girl])
                    temp_alls.append(g_invi[girl])
                    updateStaff(g_invi[girl])
                    work[g_invi[girl]][dates[dates_number]]=room
                    girl += 1
                    if girl == len(g_invi):
                        leftshift(g_invi)
                        girl = 0

                elif room in singles:

                    if room in boys_roomAllocation:
                        if b_invi[staffIndex] not in temp_alls:
                            boys_roomAllocation[room].append(b_invi[staffIndex])
                            updateStaff(b_invi[staffIndex])
                            temp_alls.append(b_invi[staffIndex])
                            work[b_invi[staffIndex]][dates[dates_number]]=room
                            staffIndex += 1
                            if staffIndex == len(b_invi):
                                leftshift(b_invi)
                                staffIndex = 0
                        else:
                            boys_roomAllocation[room].append(g_invi[girl])
                            temp_alls.append(g_invi[girl])
                            updateStaff(g_invi[girl])
                            work[g_invi[girl]][dates[dates_number]]=room
                            girl += 1
                            if girl == len(g_invi):
                                leftshift(g_invi)
                                girl = 0

                    
                    elif room in girls_roomAllocation:
                        girls_roomAllocation[room].append(g_invi[girl])
                        temp_alls.append(g_invi[girl])
                        updateStaff(g_invi[girl])
                        work[g_invi[girl]][dates[dates_number]]=room
                        girl += 1
                        if girl == len(g_invi):
                            leftshift(g_invi)
                            girl = 0
                    
                elif room in boys_roomAllocation:
                    if b_invi[staffIndex] not in temp_alls:
                        boys_roomAllocation[room].append(b_invi[staffIndex])
                        temp_alls.append(b_invi[staffIndex])
                        updateStaff(b_invi[staffIndex])
                        work[b_invi[staffIndex]][dates[dates_number]]=room
                        staffIndex += 1
                        if staffIndex == len(b_invi):
                            leftshift(b_invi)
                            staffIndex = 0
                    else:
                        boys_roomAllocation[room].append(g_invi[girl])
                        temp_alls.append(g_invi[girl])
                        updateStaff(g_invi[girl])
                        work[g_invi[girl]][dates[dates_number]]=room
                        girl += 1
                        if girl == len(g_invi):
                            leftshift(g_invi)
                            girl = 0


            for room in all_roomsAlloc:
                
                if room in girls_roomAllocation and room not in singles and g_invi[girl] not in temp_alls:
                    girls_roomAllocation[room].append(g_invi[girl])
                    temp_alls.append(g_invi[girl])
                    updateStaff(g_invi[girl])
                    work[g_invi[girl]][dates[dates_number]]=room
                    girl += 1
                    if girl == len(g_invi):
                        leftshift(g_invi)
                        girl = 0

                elif room in singles:
                    continue
                    
                elif room in boys_roomAllocation:
                    if b_invi[staffIndex] not in temp_alls:
                        boys_roomAllocation[room].append(b_invi[staffIndex])
                        temp_alls.append(b_invi[staffIndex])
                        updateStaff(b_invi[staffIndex])
                        work[b_invi[staffIndex]][dates[dates_number]]=room
                        staffIndex += 1
                        if staffIndex == len(b_invi):
                            leftshift(b_invi)
                            staffIndex = 0
                    else:
                        boys_roomAllocation[room].append(g_invi[girl])
                        temp_alls.append(g_invi[girl])
                        updateStaff(g_invi[girl])
                        work[g_invi[girl]][dates[dates_number]]=room
                        girl += 1
                        if girl == len(g_invi):
                            leftshift(g_invi)
                            girl = 0
            
                '''

                Code for if only one floor (test)

                if(gender[invigilators[staffIndex]] == 'F' and k<len(girls_rooms)):
                    girls_roomAllocation[f'S{k+1}'].append(invigilators[staffIndex])
                    updateStaff(invigilators[staffIndex])
                    #To add each assigned classroom to staffs
                    work[invigilators[staffIndex]][dates[dates_number]]=f'S{k+1}'
                    staffIndex += 1
                    if(len(girls_roomAllocation[f'S{k+1}'])==2):
                        k+=1
                
                else:
                    if(f"S{l+1}" in not_needed):
                        l+=1
                        continue
                        
                    boys_roomAllocation[f'S{l+1}'].append(invigilators[staffIndex])
                    #To add each assigned classroom to staffs
                    work[invigilators[staffIndex]][dates[dates_number]]=f'S{l+1}'

                    updateStaff(invigilators[staffIndex])
                    staffIndex += 1
                    if(f"S{l+1}" == "S24"):
                        l+=1
                    elif(f"S{l+1}"=="S26"):
                        break
                    else:
                        if(len(boys_roomAllocation[f'S{l+1}'])==2):
                            l+=1

                '''

            document = Document()
            
            # logo_path = 'CITL.png'
            # top_path = 'Top.png'

            # paragraph = document.add_paragraph()
            # run = paragraph.add_run()
            # run.add_picture(logo_path,width=Cm(10))
            # run_2 = paragraph.add_run()
            # run_2.add_picture(top_path,width=Cm(5))
            
            header = document.add_table(rows=2, cols=1)
            header.style = 'Table Grid'
            hdr_cells = header.rows[0].cells
            hdr_cells[0].text = exam
            hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.6)
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells = header.rows[1].cells
            hdr_cells[0].text = f'\n       Exam duty list           Date={dates[dates_number]}\n'
            hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.4)
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True

            
            dates_number+=1
            
            table = document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'

            for cell in table.columns[0].cells:
                cell.width = Cm(1)
            for cell in table.columns[1].cells:
                cell.width = Cm(8)
            for cell in table.columns[2].cells:
                cell.width = Cm(6)
            for cell in table.columns[3].cells:
                cell.width = Cm(6)
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '\nHall\n'
            hdr_cells[1].text = '\nName of the Faculty\n'
            hdr_cells[2].text = '\nTime\n'
            hdr_cells[3].text = '\nSignature\n'
            hdr_cells[0].paragraphs[0].runs[0].font.bold = True
            hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            hdr_cells[1].paragraphs[0].runs[0].font.bold = True
            hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[1].paragraphs[0].runs[0].font.name = 'Arial'
            hdr_cells[2].paragraphs[0].runs[0].font.bold = True
            hdr_cells[2].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[2].paragraphs[0].runs[0].font.name = 'Arial'
            hdr_cells[3].paragraphs[0].runs[0].font.bold = True
            hdr_cells[3].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[3].paragraphs[0].runs[0].font.name = 'Arial'



            temp_cell = 0

            #for girls
            for room, staffs in girls_roomAllocation.items():
                if room in singles:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'\n{room}\n'
                    for i in staffs:
                        row_cells[1].text = f'\n{i}\n'
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'

                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'\n{room}\n'
                    row_cells[1].text = '\n'.join(staffs)

                    '''
                    internal = document.add_table(rows=2, cols=4)
                    internal.style = 'Table Grid'
                    row_cells = internal.add_row().cells
                    row_cells[0].text = f'\n{staffs[0]}'
                    hdr_cells[1].text = f'\n{staffs[1]}'
                    '''





            
            #for boys
            for room, staffs in boys_roomAllocation.items():
                if room in singles:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'\n{room}\n'
                    for i in staffs:
                        row_cells[1].text = f'\n{i}\n'
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'

                else:
                    row_cells = table.add_row().cells
                    row_cells[0].text = f'\n{room}\n'
                    row_cells[1].text = '\n'.join(staffs)
                    row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
                    row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'


            '''

            Squad members table

            squad = document.add_table(rows=4, cols=4)
            squad.style = 'Table Grid'
            
            hdr_cells = squad.rows[0].cells
            hdr_cells[0].text = 'Hall'
            hdr_cells[1].text = 'Name of the Faculty'
            hdr_cells[2].text = 'Time'
            hdr_cells[3].text = 'Signature'

            '''
            
            document.save(f"Day_{j+1}_Room_Allocation.docx")
            docx2pdf.convert(f"Day_{j+1}_Room_Allocation.docx")
            os.remove(f"Day_{j+1}_Room_Allocation.docx")

            day_list.append(f"Day_{j+1}_Room_Allocation.pdf")
        

        merger = PdfMerger()
        for pdf in day_list:
            merger.append(pdf)
        merger.write("merged_days.pdf")

        # for pdf in day_list:
        #     os.remove(pdf)



        
        document = Document()

        # logo_path = 'CITL.png'
        # top_path = 'Top.png'

        # paragraph = document.add_paragraph()
        # run = paragraph.add_run()
        # run.add_picture(logo_path,width=Cm(10))
        # run_2 = paragraph.add_run()
        # run_2.add_picture(top_path,width=Cm(5))

        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nStaff Name\n'
        hdr_cells[1].text = '\nWork Count\n'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.4)
        hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells[1].paragraphs[0].runs[0].font.bold = True
        hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[1].paragraphs[0].runs[0].font.size = Cm(0.4)
        hdr_cells[1].paragraphs[0].runs[0].font.name = 'Arial'



        for staff, count in Invigilator_work_count.items():
            row_cells = table.add_row().cells
            row_cells[0].text = f'\n{staff}\n'
            row_cells[1].text = f'\n{str(count)}\n'
            row_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
            row_cells[1].paragraphs[0].runs[0].font.name = 'Arial'

        document.save("Invigilator_Work_Count.docx")
        docx2pdf.convert("Invigilator_Work_Count.docx")
        os.remove("Invigilator_Work_Count.docx")

        document = Document()

        # logo_path = 'CITL.png'
        # top_path = 'Top.png'

        # paragraph = document.add_paragraph()
        # run = paragraph.add_run()
        # run.add_picture(logo_path,width=Cm(10))
        # run_2 = paragraph.add_run()
        # run_2.add_picture(top_path,width=Cm(5))
        
        header = document.add_table(rows=2, cols=1)
        header.style = 'Table Grid'
        hdr_cells = header.rows[0].cells
        hdr_cells[0].text = '''
        Office of the Controller of Examination
        '''
        hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.4)
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'
        hdr_cells = header.rows[1].cells
        hdr_cells[0].text = '''
        Internal Assessment I AUGUST-2023-DUTY LIST
        '''
        hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.4)
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'


        table = document.add_table(rows=1, cols=len(dates)+1)
        for m in range(len(dates)+1):
            for cell in table.columns[m].cells:
                cell.width = Inches(6)
        
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '\nStaff Name\n'
        hdr_cells[0].paragraphs[0].runs[0].font.bold = True
        hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[0].paragraphs[0].runs[0].font.size = Cm(0.4)
        hdr_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

        for a in range(len(dates)):
            hdr_cells[a+1].text = f'\n{dates[a]}\n'
            hdr_cells[a+1].paragraphs[0].runs[0].font.bold = True
            hdr_cells[a+1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[a+1].paragraphs[0].runs[0].font.size = Cm(0.4)
            hdr_cells[a+1].paragraphs[0].runs[0].font.name = 'Arial'

        for staff, days in work.items():
            row_cells = table.add_row().cells
            row_cells[0].text = staff
            row_cells[0].paragraphs[0].runs[0].font.name = 'Arial'

            col=1
            for a in days.values():
                row_cells[col].text = a
                col+=1

        # print(work)
        document.save("Invigilator_Shedule.docx")
        docx2pdf.convert("Invigilator_Shedule.docx")
        os.remove("Invigilator_Shedule.docx")



    def updateStaff(staff):
        Invigilator_work_count[staff] += 1


    # Required Datas 
    dataset = pd.read_csv('Newstaffs.csv', encoding= 'unicode_escape')

    invigilators = mselected_faculty+fselected_faculty

    b_invi = mselected_faculty
    g_invi = fselected_faculty

    random.shuffle(b_invi)
    random.shuffle(g_invi)


    #to get the gender


    #Boys and Girls Input List
    '''
    First__rooms_available=[f'F{i}' for i in range(27)]
    girl_needed_rooms=[]
    boy_needed_rooms=[]
    '''

    '''All Needed rooms'''

    all_rooms = rooms
    '''Singles staff rooms'''

    singles = single 
    girls_rooms={}
    '''Boys rooms dictionary'''
    for i in girls:
        girls_rooms[i]=[]
    gir=set(girls)
    allr=set(rooms)
    byr=allr-gir
    boys_rooms={}
    for i in byr:
        boys_rooms[i]=[]
    Invigilator_work_count = {i:0 for i in invigilators}



    work={i:{j:"-" for j in dates} for i in invigilators}

    # Shuffle Staffs
    random.shuffle(invigilators)

    # Rooms needed each day (can be extended)
    roomsPerDay = [len(rooms) for i in range(tot)]   

    generateResult(roomsPerDay,girls)

# for pdf in day_list:
#      os.remove(pdf)